import os
import json
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document

from app.tools.base import BaseTool


class CSVTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="csv",
            description="Read and write CSV files",
            schema={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "write"],
                        "description": "Action to perform"
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to CSV file"
                    },
                    "data": {
                        "type": "array",
                        "description": "Data to write (required for write action)"
                    },
                    "delimiter": {
                        "type": "string",
                        "default": ",",
                        "description": "CSV delimiter"
                    },
                    "encoding": {
                        "type": "string",
                        "default": "utf-8",
                        "description": "File encoding"
                    }
                },
                "required": ["action", "file_path"]
            }
        )

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        action = parameters["action"]
        file_path = parameters["file_path"]

        try:
            if action == "read":
                return self._read_csv(file_path, parameters)
            elif action == "write":
                return self._write_csv(file_path, parameters)
            else:
                raise ValueError(f"Unknown action: {action}")

        except Exception as e:
            raise Exception(f"CSV file error: {e}")

    def _read_csv(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        delimiter = parameters.get("delimiter", ",")
        encoding = parameters.get("encoding", "utf-8")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        df = pd.read_csv(file_path, delimiter=delimiter, encoding=encoding)

        return {
            "file_path": file_path,
            "data": df.to_dict('records'),
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "encoding": encoding,
            "delimiter": delimiter
        }

    def _write_csv(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        data = parameters["data"]
        delimiter = parameters.get("delimiter", ",")
        encoding = parameters.get("encoding", "utf-8")

        if not data:
            raise ValueError("No data provided for writing")

        df = pd.DataFrame(data)

        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        df.to_csv(file_path, index=False, sep=delimiter, encoding=encoding)

        return {
            "file_path": file_path,
            "rows_written": len(df),
            "columns": df.columns.tolist(),
            "encoding": encoding,
            "delimiter": delimiter
        }


class JSONTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="json",
            description="Read and write JSON files",
            schema={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "write"],
                        "description": "Action to perform"
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to JSON file"
                    },
                    "data": {
                        "type": "object",
                        "description": "Data to write (required for write action)"
                    },
                    "encoding": {
                        "type": "string",
                        "default": "utf-8",
                        "description": "File encoding"
                    },
                    "indent": {
                        "type": "integer",
                        "default": 2,
                        "description": "JSON indentation"
                    }
                },
                "required": ["action", "file_path"]
            }
        )

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        action = parameters["action"]
        file_path = parameters["file_path"]

        try:
            if action == "read":
                return self._read_json(file_path, parameters)
            elif action == "write":
                return self._write_json(file_path, parameters)
            else:
                raise ValueError(f"Unknown action: {action}")

        except Exception as e:
            raise Exception(f"JSON file error: {e}")

    def _read_json(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        encoding = parameters.get("encoding", "utf-8")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        with open(file_path, 'r', encoding=encoding) as f:
            data = json.load(f)

        return {
            "file_path": file_path,
            "data": data,
            "encoding": encoding
        }

    def _write_json(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        data = parameters["data"]
        encoding = parameters.get("encoding", "utf-8")
        indent = parameters.get("indent", 2)

        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        with open(file_path, 'w', encoding=encoding) as f:
            json.dump(data, f, indent=indent, ensure_ascii=False)

        return {
            "file_path": file_path,
            "encoding": encoding,
            "indent": indent
        }


class FileListTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="file_list",
            description="List files in a directory",
            schema={
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "Directory path to list"
                    },
                    "pattern": {
                        "type": "string",
                        "description": "File pattern (e.g., '*.csv', '*.json')"
                    },
                    "recursive": {
                        "type": "boolean",
                        "default": False,
                        "description": "Search recursively"
                    }
                },
                "required": ["directory"]
            }
        )

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        directory = parameters["directory"]
        pattern = parameters.get("pattern", "*")
        recursive = parameters.get("recursive", False)

        try:
            if not os.path.exists(directory):
                raise FileNotFoundError(f"Directory not found: {directory}")

            if not os.path.isdir(directory):
                raise NotADirectoryError(f"Not a directory: {directory}")

            if recursive:
                files = []
                for root, dirs, filenames in os.walk(directory):
                    for filename in filenames:
                        if self._matches_pattern(filename, pattern):
                            full_path = os.path.join(root, filename)
                            files.append({
                                "path": full_path,
                                "name": filename,
                                "size": os.path.getsize(full_path),
                                "modified": os.path.getmtime(full_path)
                            })
            else:
                files = []
                for filename in os.listdir(directory):
                    if self._matches_pattern(filename, pattern):
                        full_path = os.path.join(directory, filename)
                        if os.path.isfile(full_path):
                            files.append({
                                "path": full_path,
                                "name": filename,
                                "size": os.path.getsize(full_path),
                                "modified": os.path.getmtime(full_path)
                            })

            return {
                "directory": directory,
                "pattern": pattern,
                "recursive": recursive,
                "files": files,
                "count": len(files)
            }

        except Exception as e:
            raise Exception(f"File listing error: {e}")

    def _matches_pattern(self, filename: str, pattern: str) -> bool:
        import fnmatch
        return fnmatch.fnmatch(filename, pattern)


class DocxTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="docx",
            description="Read, write, and append DOCX files",
            schema={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "write", "append"],
                        "description": "Action to perform",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to DOCX file",
                    },
                    "paragraphs": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Paragraphs to write or append",
                    },
                    "overwrite": {
                        "type": "boolean",
                        "default": True,
                        "description": "Allow overwriting existing file when writing",
                    },
                },
                "required": ["action", "file_path"],
            },
        )

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        action = parameters["action"]
        file_path = parameters["file_path"]

        try:
            if action == "read":
                return self._read_docx(file_path)
            if action == "write":
                return self._write_docx(file_path, parameters)
            if action == "append":
                return self._append_docx(file_path, parameters)
            raise ValueError(f"Unknown action: {action}")
        except Exception as exc:  # pragma: no cover - surfaced to caller
            raise Exception(f"DOCX file error: {exc}")

    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        document = Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        word_count = sum(len(p.split()) for p in paragraphs)

        return {
            "file_path": file_path,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "word_count": word_count,
        }

    def _write_docx(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        paragraphs: Optional[List[str]] = parameters.get("paragraphs")
        overwrite: bool = parameters.get("overwrite", True)

        if not paragraphs:
            raise ValueError("No paragraphs provided for writing")

        path = Path(file_path)
        if path.exists() and not overwrite:
            raise FileExistsError(f"File already exists: {file_path}")

        if path.parent:
            path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(file_path)

        return {
            "file_path": file_path,
            "paragraphs_written": len(paragraphs),
            "overwrite": overwrite,
        }

    def _append_docx(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        paragraphs: Optional[List[str]] = parameters.get("paragraphs")

        if not paragraphs:
            raise ValueError("No paragraphs provided for appending")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        document = Document(file_path)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(file_path)

        return {
            "file_path": file_path,
            "paragraphs_appended": len(paragraphs),
        }


class SpreadsheetTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="spreadsheet",
            description="Read and write Excel spreadsheets (.xlsx/.xls)",
            schema={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["read", "write"],
                        "description": "Action to perform",
                    },
                    "file_path": {
                        "type": "string",
                        "description": "Path to spreadsheet file",
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "Worksheet name (default first sheet)",
                    },
                    "data": {
                        "type": "array",
                        "description": "Rows to write (list of dicts or lists)",
                    },
                    "columns": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Column names when providing list-based rows",
                    },
                    "overwrite": {
                        "type": "boolean",
                        "default": True,
                        "description": "Allow overwriting an existing file",
                    },
                },
                "required": ["action", "file_path"],
            },
        )

    def execute(self, parameters: Dict[str, Any]) -> Dict[str, Any]:
        action = parameters["action"]
        file_path = parameters["file_path"]

        try:
            if action == "read":
                return self._read_spreadsheet(file_path, parameters)
            if action == "write":
                return self._write_spreadsheet(file_path, parameters)
            raise ValueError(f"Unknown action: {action}")
        except Exception as exc:  # pragma: no cover - surfaced to caller
            raise Exception(f"Spreadsheet error: {exc}")

    def _read_spreadsheet(self, file_path: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = Path(file_path).suffix.lower()
        if suffix not in {".xlsx", ".xls", ".xlsm"}:
            raise ValueError("Unsupported spreadsheet format; use .xlsx or .xls")

        engine = "openpyxl" if suffix in {".xlsx", ".xlsm"} else None
        sheet_name: Optional[str] = parameters.get("sheet_name")

        excel_file = pd.ExcelFile(file_path, engine=engine)
        resolved_sheet = sheet_name if sheet_name is not None else excel_file.sheet_names[0]

        df = excel_file.parse(resolved_sheet)

        return {
            "file_path": file_path,
            "sheet_name": resolved_sheet,
            "data": df.to_dict("records"),
            "columns": df.columns.tolist(),
            "row_count": len(df),
        }

    def _write_spreadsheet(
        self, file_path: str, parameters: Dict[str, Any]
    ) -> Dict[str, Any]:
        if "data" not in parameters:
            raise ValueError("No data provided for writing")

        data = parameters["data"]
        columns: Optional[List[str]] = parameters.get("columns")
        sheet_name: str = parameters.get("sheet_name", "Sheet1")
        overwrite: bool = parameters.get("overwrite", True)

        path = Path(file_path)
        if path.exists() and not overwrite:
            raise FileExistsError(f"File already exists: {file_path}")

        if path.parent:
            path.parent.mkdir(parents=True, exist_ok=True)

        df = self._data_to_dataframe(data, columns)
        df.to_excel(file_path, index=False, sheet_name=sheet_name, engine="openpyxl")

        return {
            "file_path": file_path,
            "sheet_name": sheet_name,
            "rows_written": len(df),
            "columns": df.columns.tolist(),
            "overwrite": overwrite,
        }

    def _data_to_dataframe(
        self, data: Any, columns: Optional[List[str]]
    ) -> pd.DataFrame:
        if isinstance(data, list) and (not data or isinstance(data[0], dict)):
            return pd.DataFrame(data)
        if isinstance(data, list) and isinstance(data[0], (list, tuple)):
            return pd.DataFrame(data, columns=columns)
        raise ValueError("Data must be a list of dictionaries or list of lists")
